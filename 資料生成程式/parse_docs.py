# -*- coding: utf-8 -*-
"""
法案文件內容解析器 (v8 - 自動化批次版)

功能：
1. 自動掃描來源資料夾中的所有子資料夾 (例如 "2025_08", "2025_09")。
2. 針對每一個子資料夾，讀取其中所有的 .docx 檔案。
3. 分別提取文件中的「段落文字」與「表格內容」。
4. 將每個子資料夾的結構化結果，儲存到以該子資料夾名稱命名的 .json 檔案中。
   (例如，來源 "2025_08" 的資料會存成 "structured_texts_2025_08.json")
"""
import os
import docx
import json

# --- 1. 設定來源與輸出根目錄 ---
input_folder = r"C:\Users\weiwe\Desktop\legislative_ai_web\storage\docx"
output_folder = r"C:\Users\weiwe\Desktop\legislative_ai_web\storage\docx_output"

# --- 檢查路徑並建立輸出目錄 ---
if not os.path.exists(input_folder):
    print(f"❌ 錯誤：找不到來源資料夾 '{input_folder}'。請確認路徑是否正確。")
    exit()
os.makedirs(output_folder, exist_ok=True)

# --- 2. 獲取所有要處理的子資料夾 ---
try:
    # 篩選出 input_folder 中所有「資料夾」項目
    subdirectories = [d for d in os.listdir(input_folder) if os.path.isdir(os.path.join(input_folder, d))]
    if not subdirectories:
        print(f"ℹ️  在 '{input_folder}' 中沒有找到任何子資料夾可供處理。")
        exit()
    print(f"🔍 發現 {len(subdirectories)} 個子資料夾，準備開始處理...")
except Exception as e:
    print(f"❌ 掃描子資料夾時發生錯誤: {e}")
    exit()

# --- 3. 遍歷每一個子資料夾並處理其中的檔案 ---
for subdir_name in subdirectories:
    print(f"\n{'='*60}")
    print(f"▶️  開始處理子資料夾: {subdir_name}")
    
    current_input_path = os.path.join(input_folder, subdir_name)
    all_files = os.listdir(current_input_path)
    
    # 針對目前子資料夾的資料，初始化一個儲存字典
    extracted_data_for_subdir = {}
    
    docx_files_in_subdir = [f for f in all_files if f.lower().endswith('.docx')]
    if not docx_files_in_subdir:
        print(f"ℹ️  在 '{subdir_name}' 中沒有找到 .docx 檔案，跳過此資料夾。")
        continue

    print(f"   共有 {len(docx_files_in_subdir)} 個 .docx 檔案待處理。")

    for index, filename in enumerate(docx_files_in_subdir):
        file_path = os.path.join(current_input_path, filename)
        
        print("-" * 50)
        print(f"   處理進度 ({subdir_name}): {index + 1} / {len(docx_files_in_subdir)}")
        print(f"   正在讀取檔案: {filename}")

        try:
            doc = docx.Document(file_path)
            
            # 準備存放該檔案的結構化內容
            parsed_content = {
                "paragraphs": [para.text.strip() for para in doc.paragraphs if para.text.strip()],
                "tables": []
            }

            # 專門處理表格
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                parsed_content["tables"].append(table_data)
            
            # 將這個檔案的內容存入該子資料夾的集合中
            extracted_data_for_subdir[filename] = parsed_content
            print("   ✅ 檔案結構化讀取成功！")

        except Exception as e:
            print(f"   ❌ 讀取檔案時發生錯誤: {filename}")
            print(f"      錯誤原因: {e}")

    # --- 4. 將當前子資料夾的結果儲存成一個 JSON 檔案 ---
    if not extracted_data_for_subdir:
        print(f"ℹ️  '{subdir_name}' 中沒有成功處理任何檔案。")
        continue
        
    # 根據子資料夾名稱動態產生輸出檔名
    output_json_filename = f"structured_texts_{subdir_name}.json"
    output_json_path = os.path.join(output_folder, output_json_filename)

    try:
        with open(output_json_path, 'w', encoding='utf-8') as f:
            json.dump(extracted_data_for_subdir, f, ensure_ascii=False, indent=2)
        print(f"✅ 已將 '{subdir_name}' 的結構化內容，成功儲存至 '{output_json_path}'")
    except Exception as e:
        print(f"❌ 儲存 JSON 檔案 '{output_json_path}' 時發生錯誤: {e}")

print(f"\n{'='*60}")
print("\n🎉🎉🎉 所有子資料夾的檔案轉換任務已全部完成！ 🎉🎉🎉")
